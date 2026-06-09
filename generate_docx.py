# -*- coding: utf-8 -*-
import os
import sys
import subprocess

# 自动安装 python-docx 库
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    except Exception as e:
        print(f"Error installing python-docx via pip: {e}. Trying to install via pip3...")
        subprocess.check_call(["pip3", "install", "python-docx"])

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_header_footer(doc):
    """
    为文档添加页眉页脚（禧语安卓 APP V1.0）
    """
    for section in doc.sections:
        # 设置页眉
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("禧语安卓 APP V1.0 —— 软件说明书")
        hrun.font.name = 'SimSun'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(128, 128, 128)
        
        # 设置页脚
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Page ")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(9)
        frun.font.color.rgb = RGBColor(128, 128, 128)
        
        # 通过 XML 添加页码字段
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        fp._p.append(fldSimple)

def generate_document():
    doc = Document()
    
    # 页面边距设置
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        
    # 全局默认字体样式（微软正黑/宋体）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5) # 五号字
    font.color.rgb = RGBColor(51, 51, 51)
    
    # --- 封面 ---
    p_space_before = doc.add_paragraph()
    p_space_before.paragraph_format.space_before = Pt(120)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("禧语安卓 APP")
    run_title.font.name = 'Microsoft YaHei'
    run_title.font.size = Pt(30)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 82, 118)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(150)
    run_sub = p_sub.add_run("V1.0\n\n软 件 说 明 书")
    run_sub.font.name = 'Microsoft YaHei'
    run_sub.font.size = Pt(20)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(100, 110, 120)
    
    # 封面落款表格
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headers = ["开发单位：", "著作权人：", "发布日期："]
    values = ["项目研发团队", "禧语软件团队", "2026年06月"]
    
    for i in range(3):
        cell_lbl = table.cell(i, 0)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_lbl = p_lbl.add_run(headers[i])
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(11)
        r_lbl.font.name = 'Microsoft YaHei'
        
        cell_val = table.cell(i, 1)
        p_val = cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r_val = p_val.add_run(values[i])
        r_val.font.size = Pt(11)
        r_val.font.name = 'Microsoft YaHei'
        
    doc.add_page_break()
    
    # 添加页眉页脚
    add_header_footer(doc)
    
    # ------------------ 正文：直接开始编写页面和功能 ------------------
    
    # 1. 登录与注册页面
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r_h1 = p_h1.add_run("1. 登录与注册页面")
    r_h1.font.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(26, 82, 118)
    
    p_sub1 = doc.add_paragraph()
    r_sub1 = p_sub1.add_run("（一）功能描述")
    r_sub1.font.bold = True
    r_sub1.font.size = Pt(11)
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.line_spacing = 1.25
    p_desc.add_run(
        "1. 用户协议与隐私授权：用户首次打开软件进入登录界面，系统会自动唤起「用户协议与隐私政策确认弹窗」。告知用户该客户端收集必要设备信息及手机号的条款，仅在用户点击「同意」后才能继续输入验证码登录，未授权则安全退出软件。\n"
        "2. 手机验证码快捷登录：支持输入 11 位手机号码，点击获取短信验证码。输入收到的验证码后，点击「登录」即可一键完成登录或新账户的自动注册，过程简单高效。\n"
        "3. 自动登录功能：首次成功登录并绑定账号后，当下一次启动禧语时，系统将自动验证身份并直接跳过登录步骤进入消息主页，免去用户重复输入验证码的繁琐操作。"
    )
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img_desc = p_img.add_run("\n【 🔍 截图占位：请在此处插入隐私授权弹窗与手机验证码登录页面截图 】\n图 1 - 登录与注册页面\n")
    run_img_desc.font.italic = True
    run_img_desc.font.size = Pt(9.5)
    run_img_desc.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_page_break()
    
    # 2. 消息列表页面
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r_h1 = p_h1.add_run("2. 消息列表页面")
    r_h1.font.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(26, 82, 118)
    
    p_sub1 = doc.add_paragraph()
    r_sub1 = p_sub1.add_run("（一）功能描述")
    r_sub1.font.bold = True
    r_sub1.font.size = Pt(11)
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.line_spacing = 1.25
    p_desc.add_run(
        "1. 消息会话展现：登录成功后默认选中该列表，集中显示当前用户拥有的所有单聊、群聊会话。会话列表项包含聊天对象头像、昵称、最后一条消息简览、接收时间以及未读消息红点数字。\n"
        "2. 会话管理控制：用户可以长按会话列表项弹出快捷管理菜单。支持「消息置顶」（将重要好友或群置于列表最顶部优先查阅）、「免打扰模式」（屏蔽会话的新消息强提醒，保持消息静默接收）以及「删除会话」（清除此会话并清空本地的历史消息记录）。"
    )
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img_desc = p_img.add_run("\n【 🔍 截图占位：请在此处插入消息列表主界面、会话长按快捷菜单截图 】\n图 2 - 消息列表页面\n")
    run_img_desc.font.italic = True
    run_img_desc.font.size = Pt(9.5)
    run_img_desc.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_page_break()
    
    # 3. 聊天会话页面
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r_h1 = p_h1.add_run("3. 聊天会话页面")
    r_h1.font.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(26, 82, 118)
    
    p_sub1 = doc.add_paragraph()
    r_sub1 = p_sub1.add_run("（一）功能描述")
    r_sub1.font.bold = True
    r_sub1.font.size = Pt(11)
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.line_spacing = 1.25
    p_desc.add_run(
        "点击消息列表中任一好友或群组会话，即可进入专属聊天对话界面：\n"
        "1. 发送文字与表情：底栏输入框可输入纯文字消息，右侧支持快捷唤起系统动画表情面板，一键发送生动可爱的表情。\n"
        "2. 发送多媒体消息：支持多媒体发送功能，用户点击底栏加号，可选取本地相册中已有的图片、直接调用相机拍摄照片或视频，并将它们发送给聊天对象。\n"
        "3. 发送语音消息：在不便打字时，用户长按界面底部的语音按钮进行说话录音，松手后即可将该段语音一键发送给对方，便于高效沟通。\n"
        "4. 发送文件：支持发送多类型本地文件。点击多功能展开面板的“文件”按钮，可浏览并选择手机本地存储中的文档、压缩包、安装包等文件发送给对方。\n"
        "5. 送达与已读未读状态：所有发出的消息旁边会显示实时的投递与查阅状态，单聊会话中能直接显示「已送达」以及「已读/未读」，确保发送者一目了然。"
    )
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img_desc = p_img.add_run("\n【 🔍 截图占位：请在此处插入聊天对话窗口、多功能底栏工具栏、表情面板截图 】\n图 3 - 聊天会话页面\n")
    run_img_desc.font.italic = True
    run_img_desc.font.size = Pt(9.5)
    run_img_desc.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_page_break()
    
    # 4. 通讯录页面
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r_h1 = p_h1.add_run("4. 通讯录页面")
    r_h1.font.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(26, 82, 118)
    
    p_sub1 = doc.add_paragraph()
    r_sub1 = p_sub1.add_run("（一）功能描述")
    r_sub1.font.bold = True
    r_sub1.font.size = Pt(11)
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.line_spacing = 1.25
    p_desc.add_run(
        "1. 好友统一排列：展示用户拥有的所有联系人好友。列表按首字母拼音 A-Z 顺序排列，右侧配有快捷索引条，滑动即可实现极速联系人查找。\n"
        "2. 精准添加好友：点击右上角，可输入好友的 11 位手机号进行精确搜索并发送加友申请。同时支持点击「新的朋友」页面接收、查看并「同意」或「拒绝」他人的好友申请。\n"
        "3. 扫一扫功能：内置高精度二维码扫码器。用户可扫描他人提供的「我的名片」个人二维码名片进行快速好友添加申请，亦可扫描群二维码快速一步加群交流。"
    )
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img_desc = p_img.add_run("\n【 🔍 截图占位：请在此处插入通讯录列表页、添加好友搜索页、新的朋友申请列表截图 】\n图 4 - 通讯录页面\n")
    run_img_desc.font.italic = True
    run_img_desc.font.size = Pt(9.5)
    run_img_desc.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_page_break()
    
    # 5. 群组管理页面
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r_h1 = p_h1.add_run("5. 群组管理页面")
    r_h1.font.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(26, 82, 118)
    
    p_sub1 = doc.add_paragraph()
    r_sub1 = p_sub1.add_run("（一）功能描述")
    r_sub1.font.bold = True
    r_sub1.font.size = Pt(11)
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.line_spacing = 1.25
    p_desc.add_run(
        "1. 一键创建群组：用户点击会话或通讯录右上角，选择「发起群聊」，在好友列表中勾选多人点击确定，即可迅速生成多人群聊会话，开始多人实时互动沟通。\n"
        "2. 群公告与群名修改：群主和群管理员在群设置里，支持修改群聊名称，以及撰写、编辑群公告并一键向全体群员强提醒通知，避免关键讯息遗失。\n"
        "3. 群秩序管理：群主和群管理员拥有强大的群组管理权限。支持在成员列表中对特定违规成员进行「群内禁言」、一键「移出群聊」。普通成员可主动选择「退出群聊」，群主可以直接选择「解散群聊」销毁该群组。"
    )
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img_desc = p_img.add_run("\n【 🔍 截图占位：请在此处插入群设置页面、修改群名、发布群公告、成员管理页截图 】\n图 5 - 群组管理页面\n")
    run_img_desc.font.italic = True
    run_img_desc.font.size = Pt(9.5)
    run_img_desc.font.color.rgb = RGBColor(231, 76, 60)
    
    doc.add_page_break()
    
    # 6. 个人中心页面
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r_h1 = p_h1.add_run("6. 个人中心页面")
    r_h1.font.bold = True
    r_h1.font.size = Pt(14)
    r_h1.font.color.rgb = RGBColor(26, 82, 118)
    
    p_sub1 = doc.add_paragraph()
    r_sub1 = p_sub1.add_run("（一）功能描述")
    r_sub1.font.bold = True
    r_sub1.font.size = Pt(11)
    
    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.line_spacing = 1.25
    p_desc.add_run(
        "1. 个人资料维护：「我的」个人主页展示了当前登录用户的头像、昵称、个人账号名。点击进入修改页，支持上传本地相册中选定的照片作为专属头像，支持自由修改个人专属昵称。\n"
        "2. 系统设置控制：提供系统功能设置入口。支持修改登录密码，支持一键「清除本地缓存」，清空聊天过程中堆积的临时图片和多媒体音频、视频，释放手机空间。\n"
        "3. 关于我们：点击“关于我们”可以查看“禧语 V1.0”版本号，且包含明确的版权所有声明：‘Copyright © 2026 禧语 版权所有’，提供可靠、规范、可信赖的官方信息展示。"
    )
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img_desc = p_img.add_run("\n【 🔍 截图占位：请在此处插入“我的”主界面、个人信息修改界面以及关于我们页面截图 】\n图 6 - 个人中心页面\n")
    run_img_desc.font.italic = True
    run_img_desc.font.size = Pt(9.5)
    run_img_desc.font.color.rgb = RGBColor(231, 76, 60)
    
    # 保存文档
    output_docx_path = "禧语安卓APP软件说明书.docx"
    doc.save(output_docx_path)
    print(f"Document successfully saved as {output_docx_path}")

if __name__ == '__main__':
    generate_document()
